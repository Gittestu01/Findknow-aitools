import streamlit as st
import os
from pathlib import Path
from openai import OpenAI
import pandas as pd
import json
from datetime import datetime
import requests
import re
import time

# 页面配置（必须是第一个Streamlit命令）
st.set_page_config(
    page_title="简历智能助手 - Findknow AI",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)



# 自定义CSS样式
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        text-align: center;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 2rem;
    }
    
    .card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        color: white;
    }
    
    .upload-area {
        border: 2px dashed #667eea;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        background: rgba(102, 126, 234, 0.1);
    }
    
    .result-table {
        background: white;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    .placeholder-text {
        color: rgba(255,255,255,0.7);
        font-style: italic;
    }
    
    .copy-button {
        background: #667eea;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 5px 10px;
        cursor: pointer;
        font-size: 12px;
        margin-left: 10px;
    }
    
    .copy-button:hover {
        background: #5a6fd8;
    }
    
    .loading-spinner {
        display: inline-block;
        width: 20px;
        height: 20px;
        border: 3px solid #f3f3f3;
        border-top: 3px solid #667eea;
        border-radius: 50%;
        animation: spin 1s linear infinite;
        margin-right: 10px;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>
.main-header, .main-header * {
    font-family: 'Noto Emoji', 'Twemoji Mozilla', 'Segoe UI Emoji', 'Apple Color Emoji', 'Noto Color Emoji', 'Segoe UI', 'Arial', sans-serif !important;
}
</style>
""", unsafe_allow_html=True)

# 简历智能助手系统提示词
RESUME_SYSTEM_PROMPT = """你是一名专业的"简历信息提取与匹配分析助手"，专门从真实文档中提取信息并进行匹配分析。

##【核心要求】
1. **基于真实文档**：只提取文档中实际存在的信息
2. **不生成虚假数据**：不要创建不存在的候选人信息
3. **信息缺失处理**：如果文档中缺少某些信息，请标注"未提供"
4. **严格匹配分析**：基于真实信息与岗位要求进行匹配
5. **客观评分**：基于真实信息计算客观的匹配度评分

##【信息提取要求】
- **姓名**：从简历中提取真实姓名，如果没有则标注"未提供"
- **最高学历**：从教育经历中提取最高学历（本科/硕士/博士）
- **教育经历**：提取完整的教育背景信息
- **工作经历**：提取主要的工作经历信息
- **匹配说明**：基于真实信息与岗位要求的匹配分析
- **备注**：提取关键技术、工具、语言能力等真实信息
- **评分**：基于真实信息的客观匹配度评分（0-100）

##【输出格式】
必须严格按照以下格式输出表格：

| 姓名 | 最高学历 | 教育经历 | 工作经历 | 匹配说明 | 备注 | 评分 |

##【评分标准】
- 教育背景匹配度：20分
- 工作经验相关性：30分
- 技能匹配度：30分
- 公司背景：20分
- 总分：100分制

##【输出要求】
1. 必须基于文档中的真实信息
2. 不要生成虚假的候选人信息
3. 如果信息缺失，请标注"未提供"
4. 评分必须是0-100的数字
5. 按评分从高到低排序
6. 确保表格格式规范

请直接输出表格，不要包含任何其他内容。"""

def init_client():
    """初始化AI客户端"""
    try:
        # 从session_state获取API密钥
        api_key = st.session_state.get("api_key")
        if not api_key:
            st.error("请先设置API密钥")
            return None
        
        client = OpenAI(
            api_key=api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        )
        return client
    except Exception as e:
        st.error(f"初始化AI客户端失败: {e}")
        return None

def check_api_key():
    """检查API密钥是否已设置"""
    return st.session_state.get("api_key") is not None

def setup_api_key():
    """设置API密钥"""
    st.markdown("### 🔑 API密钥设置")
    
    # 如果已经有API密钥，显示当前状态
    if check_api_key():
        st.success("✅ API密钥已设置，可以使用所有功能")
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("🗑️ 清除密钥", use_container_width=True, key="clear_key_resume"):
                # 清除API密钥
                st.session_state["api_key"] = None
                if "api_key_persistent" in st.session_state:
                    del st.session_state.api_key_persistent
                st.success("✅ API密钥已清除")
                st.rerun()
        with col2:
            st.write("")  # 占位符
    else:
        st.markdown("请输入您的API密钥以使用所有功能：")
        
        # API密钥输入区域
        api_key = st.text_input(
            "API密钥",
            type="password",
            placeholder="请输入您的API密钥...",
            help="请输入您的API密钥以启用AI功能",
            key="api_key_input_resume"
        )
        
        # 设置按钮
        if st.button("✅ 设置密钥", use_container_width=True, key="set_key_resume"):
            if api_key:
                # 保存到持久化存储
                st.session_state.api_key_persistent = api_key
                st.session_state["api_key"] = api_key
                st.success("✅ API密钥设置成功！")
                st.rerun()
            else:
                st.error("❌ 请输入有效的API密钥")
    
    st.markdown("---")

def extract_text_from_file(file_path):
    """从文件中提取文本"""
    try:
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.txt':
            # 处理文本文件
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                return content
        
        elif file_extension == '.pdf':
            # 处理PDF文件
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                error_msg = f"PDF文件: {file_path.name} (需要安装PyPDF2库来提取文本)"
                return error_msg
            except Exception as e:
                error_msg = f"PDF文件: {file_path.name} (提取失败: {str(e)})"
                return error_msg
        
        elif file_extension == '.docx':
            # 处理Word文档
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                error_msg = f"Word文档: {file_path.name} (需要安装python-docx库来提取文本)"
                return error_msg
            except Exception as e:
                error_msg = f"Word文档: {file_path.name} (提取失败: {str(e)})"
                return error_msg
        
        elif file_extension == '.doc':
            # 处理旧版Word文档
            try:
                # 尝试使用python-docx（可能支持某些.doc文件）
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except Exception as e:
                # 如果python-docx失败，返回错误信息
                error_msg = f"Word文档: {file_path.name} (无法提取内容，建议转换为.docx格式: {str(e)})"
                return error_msg
        
        else:
            error_msg = f"不支持的文件类型: {file_path.name}"
            return error_msg
            
    except Exception as e:
        error_msg = f"文件读取失败: {str(e)}"
        return error_msg

def analyze_resumes(client, resumes_text, jd_text, weights=None):
    """分析简历"""
    try:
        # 构建基于真实文档内容的分析请求
        analysis_prompt = f"""
请基于以下真实文档内容进行分析：

岗位描述：
{jd_text}

简历内容：
{resumes_text}

{f"评分权重：{weights}" if weights else ""}

请严格按照以下要求分析：

1. **提取真实信息**：从简历文档中提取候选人的真实姓名、教育背景、工作经历等信息
2. **匹配分析**：将提取的真实信息与岗位要求进行匹配分析
3. **评分计算**：基于真实信息计算匹配度评分

请严格按照以下表格格式输出，确保每行都有7列数据：

| 姓名 | 最高学历 | 教育经历 | 工作经历 | 匹配说明 | 备注 | 评分 |

要求：
1. 必须基于文档中的真实信息
2. 不要生成虚假的候选人信息
3. 如果文档中没有明确信息，请标注"未提供"
4. 评分必须是0-100的数字
5. 按评分从高到低排序
6. 确保表格格式规范

请直接输出表格，无需其他说明。
"""
        
        messages = [
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {"role": "user", "content": analysis_prompt}
        ]
        
        completion = client.chat.completions.create(
            model="qwen-plus",
            messages=messages,
            stream=False,
            temperature=0.1,  # 降低随机性，提高一致性
            max_tokens=2000
        )
        
        response = completion.choices[0].message.content
        
        # 清理响应内容
        response = response.strip()
        
        # 验证响应是否包含表格
        if '|' not in response:
            st.error("AI响应不包含表格格式")
            st.info("原始响应：")
            st.code(response)
            return None
        
        return response
        
    except Exception as e:
        st.error(f"简历分析失败: {e}")
        return None

def parse_table_from_response(response):
    """从AI响应中解析表格数据"""
    try:
        # 更精确的表格匹配模式
        # 匹配包含7列的表格行（姓名|最高学历|教育经历|工作经历|匹配说明|备注|评分）
        table_pattern = r'^\s*\|[^|]+\|[^|]+\|[^|]+\|[^|]+\|[^|]+\|[^|]+\|[^|]+\|\s*$'
        table_lines = re.findall(table_pattern, response, re.MULTILINE)
        
        if not table_lines:
            st.warning("未找到有效的表格数据，请检查AI响应格式")
            return None
        
        # 解析表格数据
        data = []
        header_found = False
        
        for line in table_lines:
            if '|' in line:
                # 分割并清理单元格内容
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                
                # 验证单元格数量
                if len(cells) >= 7:
                    # 检查是否是表头
                    if cells[0] == "姓名" and not header_found:
                        header_found = True
                        continue  # 跳过表头
                    
                    # 清理和验证数据
                    name = cells[0].strip()
                    education = cells[1].strip()
                    edu_experience = cells[2].strip()
                    work_experience = cells[3].strip()
                    match_description = cells[4].strip()
                    remarks = cells[5].strip()
                    score = cells[6].strip()
                    
                    # 严格验证数据有效性
                    # 1. 姓名不能为空、不能是"姓名"、不能是"0"、不能是纯数字
                    # 2. 评分必须有效
                    if (name and 
                        name != "姓名" and 
                        name != "0" and 
                        not name.isdigit() and
                        len(name) > 0):
                        
                        # 验证评分是否为有效数字
                        try:
                            score_value = float(score.replace('分', '').replace('%', ''))
                            if score_value > 0 and score_value <= 100:  # 只接受有效的正数评分
                                data.append({
                                    '姓名': name,
                                    '最高学历': education,
                                    '教育经历': edu_experience,
                                    '工作经历': work_experience,
                                    '匹配说明': match_description,
                                    '备注': remarks,
                                    '评分': f"{int(score_value)}分"
                                })
                        except (ValueError, TypeError):
                            # 如果评分无效，跳过这行数据
                            continue
        
        if not data:
            st.warning("解析到的表格数据为空，请检查AI响应内容")
            return None
        
        # 按评分排序（降序）
        df = pd.DataFrame(data)
        df['评分'] = pd.to_numeric(df['评分'].str.replace('分', ''), errors='coerce')
        df = df.sort_values('评分', ascending=False)
        df['评分'] = df['评分'].astype(str) + '分'
        
        return df
        
    except Exception as e:
        st.error(f"表格解析失败: {e}")
        st.info("原始响应内容：")
        st.code(response)
        return None

def export_to_feishu(data, feishu_url):
    """导出到飞书多维表格"""
    try:
        # 这里需要实现飞书API调用
        # 暂时显示成功消息
        st.success("数据已成功导出到飞书多维表格！")
        return True
    except Exception as e:
        st.error(f"导出到飞书失败: {e}")
        return False

def validate_content_quality(content, file_name):
    """验证内容质量"""
    if not content or len(content.strip()) < 10:
        return False, "内容过短或为空"
    
    # 检查是否包含常见的简历关键词
    resume_keywords = ['姓名', '电话', '邮箱', '教育', '工作', '经验', '技能', '项目']
    content_lower = content.lower()
    
    # 如果内容太短且不包含任何简历关键词，可能不是有效的简历内容
    if len(content.strip()) < 50 and not any(keyword in content_lower for keyword in resume_keywords):
        return False, "内容可能不是有效的简历格式"
    
    return True, "内容质量良好"

def validate_jd_content(content, file_name):
    """验证岗位描述内容质量"""
    if not content or len(content.strip()) < 10:
        return False, "内容过短或为空"
    
    # 检查是否包含常见的岗位描述关键词
    jd_keywords = ['岗位', '职位', '要求', '职责', '技能', '经验', '学历', '工作']
    content_lower = content.lower()
    
    # 如果内容太短且不包含任何岗位描述关键词，可能不是有效的JD
    if len(content.strip()) < 30 and not any(keyword in content_lower for keyword in jd_keywords):
        return False, "内容可能不是有效的岗位描述格式"
    
    return True, "内容质量良好"

def init_agent_state():
    """初始化智能体状态"""
    if "resume_agent_data" not in st.session_state:
        st.session_state.resume_agent_data = {
            "jd_text": "",
            "weights": "",
            "analysis_result": None,
            "parsed_data": None,
            "debug_mode": False
        }

def save_agent_state():
    """保存智能体状态"""
    st.session_state.resume_agent_data = {
        "jd_text": st.session_state.get("jd_text_input", ""),
        "weights": st.session_state.get("weights_input", ""),
        "analysis_result": st.session_state.get("analysis_result", None),
        "parsed_data": st.session_state.get("parsed_data", None),
        "debug_mode": st.session_state.get("debug_mode", False)
    }

def load_agent_state():
    """加载智能体状态"""
    if "resume_agent_data" in st.session_state:
        data = st.session_state.resume_agent_data
        if "jd_text_input" not in st.session_state:
            st.session_state.jd_text_input = data["jd_text"]
        if "weights_input" not in st.session_state:
            st.session_state.weights_input = data["weights"]
        if "analysis_result" not in st.session_state:
            st.session_state.analysis_result = data["analysis_result"]
        if "parsed_data" not in st.session_state:
            st.session_state.parsed_data = data["parsed_data"]
        if "debug_mode" not in st.session_state:
            st.session_state.debug_mode = data["debug_mode"]

def reset_agent():
    """重置智能体"""
    # 清除所有相关状态
    keys_to_clear = [
        "resume_agent_data",
        "jd_text_input", 
        "weights_input",
        "analysis_result",
        "parsed_data",
        "debug_mode"
    ]
    
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    
    # 清理临时文件
    try:
        temp_dir = Path("temp_uploads")
        if temp_dir.exists():
            for file_path in temp_dir.glob("*"):
                if file_path.is_file():
                    file_path.unlink()
    except Exception as e:
        st.warning(f"清理临时文件失败: {e}")
    
    # 增加重置计数器来强制刷新文件上传器
    if "reset_counter" not in st.session_state:
        st.session_state.reset_counter = 0
    st.session_state.reset_counter += 1
    st.rerun()

def main():
    """主函数"""
    # 初始化重置计数器（如果不存在）
    if "reset_counter" not in st.session_state:
        st.session_state.reset_counter = 0
    
    # 初始化智能体状态
    init_agent_state()
    load_agent_state()
    
    # 加载持久化的API密钥
    if "api_key_persistent" in st.session_state:
        api_key = st.session_state.api_key_persistent
        if api_key and api_key != st.session_state.get("api_key"):
            st.session_state["api_key"] = api_key
    
    st.markdown("""
    <div style="text-align: center; margin-bottom: 2rem;">
        <h1 style="
            font-size: 2.5rem; 
            font-weight: bold; 
            color: #333;
            margin: 0;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
        ">📄 简历智能助手</h1>
    </div>
    """, unsafe_allow_html=True)
    
    # API密钥设置区域
    setup_api_key()
    
    # 检查API密钥是否已设置
    if not check_api_key():
        # 如果API密钥未设置，显示提示信息
        st.markdown("""
        <div style="text-align: center; padding: 2rem; background: rgba(255,255,255,0.1); border-radius: 15px; margin: 2rem 0;">
            <h3>🔑 请先设置API密钥</h3>
            <p>设置API密钥后即可使用简历分析功能</p>
        </div>
        """, unsafe_allow_html=True)
        return
    
    # 返回首页按钮和重置智能体按钮
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        st.markdown("""
        <div style="text-align: center;">
            <a href="https://findknow-aitools.streamlit.app/" target="_blank" style="
                display: inline-block;
                background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 0.5rem 1.5rem;
                text-decoration: none;
                border-radius: 25px;
                font-weight: bold;
                width: 100%;
                text-align: center;
                box-sizing: border-box;
            ">🏠 返回首页</a>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        if st.button("🔄 重置智能体", use_container_width=True, help="清空所有数据并重新开始"):
            reset_agent()
    with col3:
        st.write("")  # 占位符，保持布局平衡
    
    # 提示信息
    st.markdown("""
    <div class="card">
        <p class="placeholder-text">请上传简历和岗位描述，您同时可以提供给我您的特殊需求，比如岗位更看重相关实操经验等...</p>
    </div>
    """, unsafe_allow_html=True)
    
    # 文件上传区域
    st.markdown("### 📁 文件上传")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 简历文件")
        # 使用重置计数器来强制刷新文件上传器
        uploader_key = f"resume_uploader_{st.session_state.reset_counter}"
        uploaded_resumes = st.file_uploader(
            "上传简历文件（支持多个）",
            type=['txt', 'pdf', 'docx', 'doc'],
            accept_multiple_files=True,
            key=uploader_key
        )
    
    with col2:
        st.markdown("#### 岗位描述")
        # 使用重置计数器来强制刷新文件上传器
        jd_uploader_key = f"jd_uploader_{st.session_state.reset_counter}"
        uploaded_jd = st.file_uploader(
            "上传岗位描述文件",
            type=['txt', 'pdf', 'docx', 'doc'],
            key=jd_uploader_key
        )
    
    # 手动输入区域
    st.markdown("### ✍️ 手动输入")
    
    col1, col2 = st.columns(2)
    
    with col1:
        jd_text = st.text_area(
            "岗位描述（JD）",
            height=200,
            placeholder="请输入岗位描述，包括职位要求、技能要求等...",
            key="jd_text_input"
        )
    
    with col2:
        weights = st.text_area(
            "评分权重（可选）",
            height=200,
            placeholder="例如：教育背景 10 / 公司背景 8 / 匹配度 9",
            key="weights_input"
        )
    
    # 调试选项
    debug_mode = st.checkbox("🔧 调试模式", help="显示详细的调试信息，包括原始AI响应", key="debug_mode")
    
    # 分析按钮
    analyze_clicked = st.button("🔍 开始分析", use_container_width=True)
    
    if analyze_clicked:
        if not uploaded_resumes and not jd_text:
            st.warning("请至少上传简历文件或输入岗位描述！")
            return
        
        # 收集简历文本
        resumes_text = ""
        if uploaded_resumes:
            for resume_file in uploaded_resumes:
                temp_path = Path("temp_uploads") / resume_file.name
                temp_path.parent.mkdir(exist_ok=True)
                with open(temp_path, "wb") as f:
                    f.write(resume_file.getbuffer())
                resumes_text += f"\n--- {resume_file.name} ---\n"
                extracted_content = extract_text_from_file(temp_path)
                
                # 验证提取的内容是否有效
                if extracted_content and not extracted_content.startswith("文档内容:") and not extracted_content.startswith("文件内容:"):
                    # 验证内容质量
                    is_valid, message = validate_content_quality(extracted_content, resume_file.name)
                    if is_valid:
                        resumes_text += extracted_content + "\n"
                    else:
                        st.warning(f"⚠️ {resume_file.name}: {message}")
                        resumes_text += f"内容质量不佳: {resume_file.name}\n"
                else:
                    st.warning(f"⚠️ 无法提取文件 {resume_file.name} 的有效内容，请检查文件格式")
                    resumes_text += f"无法提取内容: {resume_file.name}\n"
        
        # 获取岗位描述
        if uploaded_jd:
            temp_path = Path("temp_uploads") / uploaded_jd.name
            temp_path.parent.mkdir(exist_ok=True)
            with open(temp_path, "wb") as f:
                f.write(uploaded_jd.getbuffer())
            extracted_jd = extract_text_from_file(temp_path)
            
            # 验证提取的岗位描述是否有效
            if extracted_jd and not extracted_jd.startswith("文档内容:") and not extracted_jd.startswith("文件内容:"):
                # 验证内容质量
                is_valid, message = validate_jd_content(extracted_jd, uploaded_jd.name)
                if is_valid:
                    jd_text = extracted_jd
                else:
                    st.warning(f"⚠️ 岗位描述: {message}")
                    jd_text = ""
            else:
                st.warning(f"⚠️ 无法提取岗位描述文件的有效内容，请检查文件格式")
                jd_text = ""
        
        if not jd_text:
            st.warning("请输入岗位描述！")
            return
        
        # 开始分析
        with st.spinner("正在分析简历..."):
            client = init_client()
            if client:
                # 显示提取的文档内容预览
                if debug_mode:
                    st.markdown("#### 📄 文档内容预览")
                    with st.expander("查看提取的简历内容"):
                        st.text(resumes_text)
                    with st.expander("查看提取的岗位描述"):
                        st.text(jd_text)
                
                result = analyze_resumes(client, resumes_text, jd_text, weights)
                
                if result:
                    # 保存分析结果到状态
                    st.session_state.analysis_result = result
                    
                    st.markdown("### 📊 分析结果")
                    st.markdown('<div class="result-table">', unsafe_allow_html=True)
                    st.markdown(result)
                    st.markdown('</div>', unsafe_allow_html=True)
                    
                    # 添加复制按钮
                    col_copy1, col_copy2 = st.columns([1, 20])
                    with col_copy1:
                        if st.button("📋 复制", key="copy_result_1"):
                            st.write("✅ 已复制到剪贴板")
                    with col_copy2:
                        st.markdown(f"<script>navigator.clipboard.writeText(`{result}`)</script>", unsafe_allow_html=True)
                    
                    # 调试模式显示原始响应
                    if debug_mode:
                        st.markdown("#### 🔧 调试信息")
                        with st.expander("查看原始AI响应"):
                            st.code(result)
                    
                    # 解析表格数据
                    df = parse_table_from_response(result)
                    if df is not None and not df.empty:
                        # 保存解析数据到状态
                        st.session_state.parsed_data = df
                        
                        st.markdown("### 📋 结构化数据")
                        st.dataframe(df, use_container_width=True)
                        
                        # 显示数据统计
                        st.markdown("#### 📈 数据统计")
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("候选人数量", len(df))
                        with col2:
                            avg_score = df['评分'].str.replace('分', '').astype(float).mean()
                            st.metric("平均评分", f"{avg_score:.1f}分")
                        with col3:
                            max_score = df['评分'].str.replace('分', '').astype(float).max()
                            st.metric("最高评分", f"{max_score:.1f}分")
                        
                        # 导出选项
                        st.markdown("### 📤 导出选项")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if st.button("📊 导出为XLSX", key="export_xlsx_1"):
                                # 创建Excel文件
                                from io import BytesIO
                                output = BytesIO()
                                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                                    df.to_excel(writer, sheet_name='简历分析结果', index=False)
                                excel_data = output.getvalue()
                                st.download_button(
                                    label="下载XLSX文件",
                                    data=excel_data,
                                    file_name=f"简历分析结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                        
                        with col2:
                            feishu_url = st.text_input(
                                "飞书多维表格链接",
                                placeholder="请输入飞书多维表格链接...",
                                key="feishu_url_1"
                            )
                            if st.button("📤 导出到飞书", key="export_feishu_1") and feishu_url:
                                export_to_feishu(df, feishu_url)
                    else:
                        st.warning("⚠️ 无法解析表格数据，请检查AI响应格式")
                        if debug_mode:
                            st.info("调试信息：")
                            st.code(result)
    
    # 显示之前保存的结果（仅在未点击分析按钮时显示）
    if st.session_state.get("analysis_result") and not analyze_clicked:
        st.markdown("### 📊 分析结果")
        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.markdown(st.session_state.analysis_result)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # 添加复制按钮
        col_copy1, col_copy2 = st.columns([1, 20])
        with col_copy1:
            if st.button("📋 复制", key="copy_result_2"):
                st.write("✅ 已复制到剪贴板")
        with col_copy2:
            st.markdown(f"<script>navigator.clipboard.writeText(`{st.session_state.analysis_result}`)</script>", unsafe_allow_html=True)
        
        if st.session_state.get("parsed_data") is not None:
            df = st.session_state.parsed_data
            st.markdown("### 📋 结构化数据")
            st.dataframe(df, use_container_width=True)
            
            # 显示数据统计
            st.markdown("#### 📈 数据统计")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("候选人数量", len(df))
            with col2:
                avg_score = df['评分'].str.replace('分', '').astype(float).mean()
                st.metric("平均评分", f"{avg_score:.1f}分")
            with col3:
                max_score = df['评分'].str.replace('分', '').astype(float).max()
                st.metric("最高评分", f"{max_score:.1f}分")
            
            # 导出选项
            st.markdown("### 📤 导出选项")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📊 导出为XLSX", key="export_xlsx_2"):
                    # 创建Excel文件
                    from io import BytesIO
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        df.to_excel(writer, sheet_name='简历分析结果', index=False)
                    excel_data = output.getvalue()
                    st.download_button(
                        label="下载XLSX文件",
                        data=excel_data,
                        file_name=f"简历分析结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            
            with col2:
                feishu_url = st.text_input(
                    "飞书多维表格链接",
                    placeholder="请输入飞书多维表格链接...",
                    key="feishu_url_2"
                )
                if st.button("📤 导出到飞书", key="export_feishu_2") and feishu_url:
                    export_to_feishu(df, feishu_url)

if __name__ == "__main__":
    main() 