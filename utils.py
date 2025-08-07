"""
Findknow AI现代工具库工具类
"""

import os
import streamlit as st
from pathlib import Path
from openai import OpenAI
import pandas as pd
import json
from datetime import datetime
import re
from config import API_CONFIG, ERROR_MESSAGES, SUCCESS_MESSAGES, UPLOAD_CONFIG

class AIClient:
    """AI客户端管理类"""
    
    def __init__(self):
        self.client = None
        self._init_client()
    
    def _init_client(self):
        """初始化AI客户端"""
        try:
            # 从session_state获取API密钥，使用持久化存储
            api_key = st.session_state.get("api_key")
            if not api_key:
                return  # 如果没有API密钥，不初始化客户端
            
            # 修复OpenAI客户端初始化，移除不兼容的参数
            self.client = OpenAI(
                api_key=api_key,
                base_url=API_CONFIG["base_url"]
            )
        except Exception as e:
            st.error(f"{ERROR_MESSAGES['client_init_failed']}: {e}")
    
    def is_available(self):
        """检查客户端是否可用"""
        return self.client is not None and st.session_state.get("api_key") is not None
    
    def update_api_key(self, api_key):
        """更新API密钥并重新初始化客户端"""
        if api_key:
            # 使用持久化存储保存API密钥
            st.session_state["api_key"] = api_key
            # 将密钥保存到持久化存储
            if "api_key_persistent" not in st.session_state:
                st.session_state.api_key_persistent = api_key
            self._init_client()
        else:
            st.session_state["api_key"] = None
            if "api_key_persistent" in st.session_state:
                del st.session_state.api_key_persistent
            self.client = None
    
    def load_persistent_api_key(self):
        """从持久化存储加载API密钥"""
        if "api_key_persistent" in st.session_state:
            api_key = st.session_state.api_key_persistent
            if api_key and api_key != st.session_state.get("api_key"):
                st.session_state["api_key"] = api_key
                self._init_client()
                return True
        return False
    
    def save_api_key_to_persistent(self, api_key):
        """保存API密钥到持久化存储"""
        if api_key:
            st.session_state.api_key_persistent = api_key
            st.session_state["api_key"] = api_key
            self._init_client()
            return True
        return False
    
    def clear_api_key(self):
        """清除API密钥"""
        st.session_state["api_key"] = None
        if "api_key_persistent" in st.session_state:
            del st.session_state.api_key_persistent
        self.client = None
    
    def chat_completion(self, messages, model="qwen-plus", stream=True):
        """执行聊天完成"""
        if not self.is_available():
            return None
        
        try:
            # 移除不兼容的stream_options参数
            completion = self.client.chat.completions.create(
                model=model,
                messages=messages,
                stream=stream
            )
            return completion
        except Exception as e:
            st.error(f"AI对话失败: {e}")
            return None
    
    def file_upload(self, file_path, purpose="file-extract"):
        """上传文件到AI服务"""
        if not self.is_available():
            return None
        
        try:
            file_object = self.client.files.create(
                file=Path(file_path), 
                purpose=purpose
            )
            return file_object
        except Exception as e:
            st.error(f"文件上传失败: {e}")
            return None
    
    def network_search(self, query):
        """联网搜索功能"""
        if not self.is_available():
            return None
        
        try:
            completion = self.client.chat.completions.create(
                model="qwen-plus",
                messages=[
                    {'role': 'system', 'content': 'You are a helpful assistant.'},
                    {'role': 'user', 'content': query}
                ],
                stream=True,
                extra_body={
                    "enable_search": True
                }
            )
            return completion
        except Exception as e:
            st.error(f"联网搜索失败: {e}")
            return None
    
    def file_understanding(self, file_path, query):
        """文件理解功能"""
        if not self.is_available():
            return None
        
        try:
            # 确保file_path是Path对象
            if isinstance(file_path, str):
                file_path = Path(file_path)
            
            # 首先上传文件
            st.info(f"正在上传文件: {file_path.name}")
            file_object = self.file_upload(file_path, purpose="file-extract")
            if not file_object:
                st.error("文件上传失败")
                return None
            
            st.info(f"文件上传成功，ID: {file_object.id}")
            
            # 使用文件理解API
            completion = self.client.chat.completions.create(
                model="qwen-plus",
                messages=[
                    {'role': 'user', 'content': query}
                ],
                stream=True,
                extra_body={
                    "file_ids": [file_object.id]
                }
            )
            return completion
        except Exception as e:
            st.error(f"文件理解失败: {e}")
            return None

class FileManager:
    """文件管理类"""
    
    def __init__(self):
        self.temp_dir = UPLOAD_CONFIG["temp_dir"]
        # 确保临时目录存在
        self.temp_dir.mkdir(exist_ok=True)
        # 设置权限（如果需要）
        try:
            self.temp_dir.chmod(0o755)
        except:
            pass  # 忽略权限设置错误
    
    def save_uploaded_file(self, uploaded_file):
        """保存上传的文件"""
        try:
            file_path = self.temp_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            return file_path
        except Exception as e:
            st.error(f"{ERROR_MESSAGES['file_upload_failed']}: {e}")
            return None
    
    def extract_text_from_file(self, file_path):
        """从文件中提取文本"""
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.txt':
                # 处理文本文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return content
            
            elif file_extension == '.pdf':
                # 处理PDF文件
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    error_msg = f"PDF文件: {file_path.name} (需要安装PyPDF2库来提取文本)"
                    return error_msg
                except Exception as e:
                    error_msg = f"PDF文件: {file_path.name} (提取失败: {str(e)})"
                    return error_msg
            
            elif file_extension == '.docx':
                # 处理Word文档
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except ImportError:
                    error_msg = f"Word文档: {file_path.name} (需要安装python-docx库来提取文本)"
                    return error_msg
                except Exception as e:
                    error_msg = f"Word文档: {file_path.name} (提取失败: {str(e)})"
                    return error_msg
            
            elif file_extension == '.doc':
                # 处理旧版Word文档 - 使用替代方法
                try:
                    # 尝试使用python-docx（可能支持某些.doc文件）
                    from docx import Document
                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except Exception as e:
                    # 如果python-docx失败，返回错误信息
                    error_msg = f"Word文档: {file_path.name} (无法提取内容，建议转换为.docx格式: {str(e)})"
                    return error_msg
            
            else:
                error_msg = f"不支持的文件类型: {file_path.name}"
                return error_msg
                
        except Exception as e:
            error_msg = f"文件读取失败: {str(e)}"
            return error_msg
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        try:
            deleted_count = 0
            for file_path in self.temp_dir.glob("*"):
                if file_path.is_file():
                    file_path.unlink()
                    deleted_count += 1
            return deleted_count
        except Exception as e:
            return 0

class DataProcessor:
    """数据处理类"""
    
    @staticmethod
    def parse_table_from_response(response):
        """从AI响应中解析表格数据"""
        try:
            # 查找表格内容
            table_pattern = r'\|.*\|.*\|.*\|.*\|.*\|.*\|.*\|'
            table_lines = re.findall(table_pattern, response, re.MULTILINE)
            
            if not table_lines:
                return None
            
            # 解析表格数据
            data = []
            for line in table_lines[1:]:  # 跳过表头
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if len(cells) >= 7:
                        data.append({
                            '姓名': cells[0],
                            '最高学历': cells[1],
                            '教育经历': cells[2],
                            '工作经历': cells[3],
                            '匹配说明': cells[4],
                            '备注': cells[5],
                            '评分': cells[6]
                        })
            
            return pd.DataFrame(data)
        except Exception as e:
            st.error(f"表格解析失败: {e}")
            return None
    
    @staticmethod
    def export_to_csv(data, filename=None):
        """导出数据为CSV"""
        try:
            if filename is None:
                filename = f"导出数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            
            csv_data = data.to_csv(index=False, encoding='utf-8-sig')
            return csv_data, filename
        except Exception as e:
            st.error(f"CSV导出失败: {e}")
            return None, None
    
    @staticmethod
    def export_to_json(data, filename=None):
        """导出数据为JSON"""
        try:
            if filename is None:
                filename = f"导出数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            
            json_data = data.to_json(orient='records', force_ascii=False, indent=2)
            return json_data, filename
        except Exception as e:
            st.error(f"JSON导出失败: {e}")
            return None, None

class StyleManager:
    """样式管理类"""
    
    @staticmethod
    def get_main_css():
        """获取主CSS样式"""
        return """
        <style>
            .main-header {
                font-size: 3rem;
                font-weight: bold;
                text-align: center;
                background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                margin-bottom: 2rem;
            }
            
            .card {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                border-radius: 15px;
                padding: 1.5rem;
                margin: 1rem 0;
                color: white;
                transition: transform 0.3s ease;
            }
            
            .card:hover {
                transform: translateY(-5px);
                box-shadow: 0 10px 25px rgba(0,0,0,0.2);
            }
            
            .chat-container {
                background: rgba(255,255,255,0.1);
                border-radius: 10px;
                padding: 1rem;
                margin: 1rem 0;
            }
            
            .stButton > button {
                background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                border-radius: 25px;
                padding: 0.5rem 1.5rem;
                font-weight: bold;
            }
            
            .stSelectbox > div > div > select {
                background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
                color: white;
            }
            
            .placeholder-text {
                color: rgba(255,255,255,0.7);
                font-style: italic;
            }
            
            .result-area {
                background: white;
                border-radius: 10px;
                padding: 1rem;
                margin: 1rem 0;
                color: #333;
            }
            
            .upload-area {
                border: 2px dashed #667eea;
                border-radius: 10px;
                padding: 2rem;
                text-align: center;
                background: rgba(102, 126, 234, 0.1);
            }
            
            /* 小功能按钮样式 */
            .small-feature-button {
                background: rgba(255, 255, 255, 0.1);
                border: 1px solid rgba(255, 255, 255, 0.2);
                border-radius: 20px;
                padding: 0.3rem 0.8rem;
                font-size: 0.8rem;
                color: #4c4c4c;
                transition: all 0.2s ease;
                margin: 0.2rem;
            }
            
            .small-feature-button:hover {
                background: rgba(255, 255, 255, 0.2);
                border-color: rgba(255, 255, 255, 0.3);
            }
            
            /* 集成式聊天容器样式 */
            .integrated-chat-container {
                background: rgba(255, 255, 255, 0.05);
                border-radius: 15px;
                padding: 1.5rem;
                margin: 1rem 0;
                border: 1px solid rgba(255, 255, 255, 0.1);
            }

            /* 文件上传区域样式 */
            .file-upload-area {
                background: rgba(255, 255, 255, 0.05);
                border-radius: 10px;
                padding: 1rem;
                margin: 0.5rem 0;
                border: 2px dashed rgba(255, 255, 255, 0.3);
            }

            /* 聊天控制按钮样式 */
            .chat-control-buttons {
                background: rgba(255, 255, 255, 0.05);
                border-radius: 10px;
                padding: 0.5rem;
                margin: 0.5rem 0;
            }

            /* 消息删除按钮样式 */
            .message-delete-btn {
                background: rgba(255, 0, 0, 0.1);
                border: 1px solid rgba(255, 0, 0, 0.3);
                border-radius: 15px;
                padding: 0.2rem 0.5rem;
                font-size: 0.7rem;
                color: #ff4444;
                transition: all 0.2s ease;
            }

            .message-delete-btn:hover {
                background: rgba(255, 0, 0, 0.2);
                border-color: rgba(255, 0, 0, 0.5);
            }

            /* 功能开关样式 */
            .stCheckbox > div > div > label {
                color: rgba(255, 255, 255, 0.8);
                font-size: 0.9rem;
            }

            /* 聊天输入框样式优化 */
            .stChatInput > div > div > div > div > div {
                background: rgba(255, 255, 255, 0.1);
                border: 1px solid rgba(255, 255, 255, 0.2);
                border-radius: 20px;
            }

            /* 功能按钮区域样式 */
            .feature-buttons {
                background: rgba(255, 255, 255, 0.05);
                border-radius: 15px;
                padding: 1rem;
                margin: 1rem 0;
                border: 1px solid rgba(255, 255, 255, 0.1);
            }

            /* 按钮样式优化 */
            .stButton > button {
                background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                border-radius: 25px;
                padding: 0.5rem 1.5rem;
                font-weight: bold;
                transition: all 0.3s ease;
            }

            .stButton > button:hover {
                transform: translateY(-2px);
                box-shadow: 0 5px 15px rgba(102, 126, 234, 0.3);
            }

            /* 文件上传器容器样式 */
            .upload-container {
                background: rgba(255, 255, 255, 0.05);
                border-radius: 10px;
                padding: 1rem;
                margin: 0.5rem 0;
                border: 1px dashed rgba(255, 255, 255, 0.2);
            }
        </style>
        """
    
    @staticmethod
    def apply_custom_style():
        """应用自定义样式"""
        st.markdown(StyleManager.get_main_css(), unsafe_allow_html=True)

class SessionManager:
    """会话管理类"""
    
    def __init__(self):
        self.file_manager = FileManager()
    
    @staticmethod
    def init_session_state():
        """初始化会话状态"""
        # 初始化基本状态
        if "messages" not in st.session_state:
            st.session_state.messages = []
        if "current_tool" not in st.session_state:
            st.session_state.current_tool = None
        if "show_search" not in st.session_state:
            st.session_state.show_search = False
        if "show_upload" not in st.session_state:
            st.session_state.show_upload = False
        if "uploaded_files" not in st.session_state:
            st.session_state.uploaded_files = []
    
    def clear_chat_history(self):
        """清除聊天历史和临时文件"""
        # 清除基本状态
        if "messages" in st.session_state:
            st.session_state.messages = []
        
        # 清除所有文件上传相关状态
        file_upload_keys = [
            "uploaded_files",
            "previous_upload_count", 
            "files_displayed",
            "file_uploader"
        ]
        for key in file_upload_keys:
            if key in st.session_state:
                del st.session_state[key]
        
        # 清除Streamlit内部的文件上传器状态
        # 这些是Streamlit内部使用的键
        internal_file_keys = [
            "file_uploader",
            "file_uploader_uploaded_files",
            "file_uploader_uploaded_files_count"
        ]
        for key in internal_file_keys:
            if key in st.session_state:
                del st.session_state[key]
        
        # 清除所有以file_uploader开头的状态
        for key in list(st.session_state.keys()):
            if key.startswith("file_uploader"):
                del st.session_state[key]
        
        # 清除输入框相关状态
        if "input_counter" in st.session_state:
            del st.session_state.input_counter
        
        # 清除功能设置状态
        if "enable_search" in st.session_state:
            del st.session_state.enable_search
        # 删除文件上传显示状态，让它在下次初始化时重新创建
        if "show_file_upload" in st.session_state:
            del st.session_state.show_file_upload
        
        # 清除所有聊天相关组件状态
        for key in list(st.session_state.keys()):
            if (key.startswith("chat_input_") or 
                key.startswith("delete_msg_") or 
                key == "chat_form" or
                key.startswith("st.")):
                del st.session_state[key]
        
        # 清除工具选择状态
        if "tool_type" in st.session_state:
            del st.session_state.tool_type
        
        # 清除重启按钮状态
        if "restart_chat" in st.session_state:
            del st.session_state.restart_chat
        
        # 清除文件上传器重置标志
        if "reset_file_uploader" in st.session_state:
            del st.session_state.reset_file_uploader
        
        # 清理临时文件
        self.file_manager.cleanup_temp_files()
    
    @staticmethod
    def delete_last_message():
        """删除最后一条消息"""
        if "messages" in st.session_state and len(st.session_state.messages) >= 2:
            st.session_state.messages = st.session_state.messages[:-2]
    
    @staticmethod
    def delete_message(index):
        """删除指定索引的消息"""
        if "messages" in st.session_state and 0 <= index < len(st.session_state.messages):
            # 删除用户消息和对应的AI回复
            if index % 2 == 0:  # 用户消息
                if index + 1 < len(st.session_state.messages):
                    st.session_state.messages.pop(index + 1)  # 删除AI回复
                st.session_state.messages.pop(index)  # 删除用户消息
            else:  # AI消息
                if index - 1 >= 0:
                    st.session_state.messages.pop(index - 1)  # 删除用户消息
                st.session_state.messages.pop(index - 1)  # 删除AI消息
    
    @staticmethod
    def add_message(role, content):
        """添加消息到会话"""
        if "messages" not in st.session_state:
            st.session_state.messages = []
        
        st.session_state.messages.append({"role": role, "content": content})
    
    @staticmethod
    def get_messages():
        """获取所有消息"""
        return st.session_state.get("messages", []) 